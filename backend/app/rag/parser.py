from __future__ import annotations

import io
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

SOURCE_RE = re.compile(r"(?:来源|原文链接)\s*[:：]\s*(https?://\S+)", re.I)
DATE_RE = re.compile(r"发布时间\s*[:：]\s*(\d{4})[-年/](\d{1,2})[-月/](\d{1,2})")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
NOISE_RE = re.compile(r"^(上一篇|下一篇|打印|关闭窗口|返回顶部)\s*[:：]?.*$")


@dataclass(slots=True)
class ParsedSection:
    text: str
    heading_path: str = ""
    page_number: int | None = None


@dataclass(slots=True)
class ParsedDocument:
    title: str
    sections: list[ParsedSection]
    source_url: str | None = None
    published_at: date | None = None


def _metadata(text: str, fallback_title: str) -> tuple[str, str | None, date | None]:
    title = fallback_title
    for line in text.splitlines():
        match = HEADING_RE.match(line.strip())
        if match:
            title = match.group(2).strip()
            break
    source = SOURCE_RE.search(text)
    date_match = DATE_RE.search(text)
    published = None
    if date_match:
        try:
            published = date(*(int(value) for value in date_match.groups()))
        except ValueError:
            published = None
    return title, source.group(1).rstrip(")。]>") if source else None, published


def _clean_line(line: str) -> str:
    stripped = line.strip()
    if SOURCE_RE.search(stripped) or DATE_RE.search(stripped) or NOISE_RE.match(stripped):
        return ""
    return re.sub(r"[ \t]+", " ", line).strip()


def _parse_markdown(text: str, fallback_title: str) -> ParsedDocument:
    title, source, published = _metadata(text, fallback_title)
    headings: list[str] = []
    sections: list[ParsedSection] = []
    buffer: list[str] = []
    last_heading = ""

    def flush() -> None:
        nonlocal buffer
        body = "\n".join(item for item in buffer if item).strip()
        if body and body != title:
            sections.append(ParsedSection(body, last_heading))
        buffer = []

    for raw in text.replace("\r\n", "\n").splitlines():
        heading = HEADING_RE.match(raw.strip())
        if heading:
            flush()
            level = len(heading.group(1))
            value = heading.group(2).strip()
            headings[level - 1 :] = [value]
            last_heading = " > ".join(headings)
            continue
        cleaned = _clean_line(raw)
        if cleaned:
            buffer.append(cleaned)
        elif buffer and buffer[-1] != "":
            buffer.append("")
    flush()
    return ParsedDocument(title=title, sections=sections, source_url=source, published_at=published)


def parse_file(path: Path, fallback_title: str | None = None) -> ParsedDocument:
    suffix = path.suffix.lower()
    fallback = (fallback_title or path.stem).strip() or path.stem
    if suffix in {".md", ".txt"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".md":
            return _parse_markdown(text, fallback)
        title, source, published = _metadata(text, fallback)
        cleaned = "\n".join(filter(None, (_clean_line(line) for line in text.splitlines())))
        return ParsedDocument(title, [ParsedSection(cleaned)], source, published)
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(path.read_bytes()))
        sections = [
            ParsedSection((page.extract_text() or "").strip(), page_number=index)
            for index, page in enumerate(reader.pages, start=1)
            if (page.extract_text() or "").strip()
        ]
        return ParsedDocument(fallback, sections)
    if suffix == ".docx":
        doc = DocxDocument(io.BytesIO(path.read_bytes()))
        text = "\n".join(
            paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()
        )
        return _parse_markdown(text, fallback)
    raise ValueError(f"不支持的文件格式: {suffix}")
