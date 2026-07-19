from __future__ import annotations

from io import BytesIO

import pytest
from app.core.config import settings
from app.core.database import SessionLocal
from app.core.security import create_access_token, hash_password
from app.models.enums import DocumentStatus, Role
from app.models.orm import Document, User
from app.rag.parser import ParsedDocument, ParsedSection
from app.services.documents import document_service
from docx import Document as DocxDocument
from httpx import AsyncClient

pytestmark = pytest.mark.anyio


def create_admin_header() -> dict[str, str]:
    with SessionLocal() as db:
        admin = User(
            name="文档管理员",
            email="documents-admin@example.com",
            password_hash=hash_password("admin-pass-123"),
            role=Role.ADMIN,
        )
        db.add(admin)
        db.commit()
        db.refresh(admin)
        token = create_access_token(admin.id, admin.role.value)
    return {"Authorization": f"Bearer {token}"}


async def test_document_http_lifecycle_and_validation(client: AsyncClient, monkeypatch) -> None:
    header = create_admin_header()
    monkeypatch.setattr(document_service, "submit", lambda _document_id, _job_id: None)

    registered = await client.post(
        "/api/auth/register",
        json={"name": "普通学生", "email": "student-docs@example.com", "password": "safe-pass-123"},
    )
    student_header = {
        "Authorization": f"Bearer {registered.json()['data']['access_token']}"
    }
    visible = await client.get("/api/documents", headers=student_header)
    assert visible.status_code == 200

    rejected = await client.post(
        "/api/documents",
        headers=header,
        files={"file": ("malware.exe", b"not allowed", "application/octet-stream")},
    )
    assert rejected.status_code == 400

    content = (
        "# 新生办事指南\n\n> 来源：https://example.edu.cn/guide\n\n新生报到时请携带录取通知书。"
    ).encode()
    uploaded = await client.post(
        "/api/documents",
        headers=header,
        files={"file": ("../guide.md", content, "text/markdown")},
        data={"title": "新生办事指南", "category": "campus_life"},
    )
    assert uploaded.status_code == 202
    document_id = uploaded.json()["data"]["document"]["id"]
    assert uploaded.json()["data"]["document"]["original_name"] == "guide.md"

    listed = await client.get("/api/documents?q=新生", headers=header)
    assert listed.status_code == 200
    assert listed.json()["data"]["total"] == 1

    detail = await client.get(f"/api/documents/{document_id}", headers=header)
    assert detail.status_code == 200
    assert detail.json()["data"]["category"] == "campus_life"

    student_detail = await client.get(
        f"/api/documents/{document_id}", headers=student_header
    )
    assert student_detail.status_code == 200
    preview = await client.get(
        f"/api/documents/{document_id}/preview?offset=2&limit=8",
        headers=student_header,
    )
    assert preview.status_code == 200
    assert preview.json()["data"]["offset"] == 2
    assert preview.json()["data"]["limit"] == 8
    assert preview.json()["data"]["total_chars"] == len(content.decode())
    assert preview.json()["data"]["has_more"] is True
    capped_preview = await client.get(
        f"/api/documents/{document_id}/preview?limit=999999",
        headers=student_header,
    )
    assert capped_preview.json()["data"]["limit"] == 50_000
    download = await client.get(
        f"/api/documents/{document_id}/download", headers=student_header
    )
    assert download.status_code == 200
    assert download.content == content
    assert "guide.md" in download.headers["content-disposition"]

    for method, path in (
        ("patch", f"/api/documents/{document_id}"),
        ("post", f"/api/documents/{document_id}/reindex"),
        ("delete", f"/api/documents/{document_id}"),
    ):
        response = await client.request(
            method,
            path,
            headers=student_header,
            json={"title": "x", "category": "x"} if method == "patch" else None,
        )
        assert response.status_code == 403

    edit_conflict = await client.patch(
        f"/api/documents/{document_id}",
        headers=header,
        json={"title": "新标题", "category": "通知", "source_url": None, "published_at": None},
    )
    assert edit_conflict.status_code == 409

    duplicate = await client.post(
        "/api/documents",
        headers=header,
        files={"file": ("copy.md", content, "text/markdown")},
    )
    assert duplicate.status_code == 409

    processing_conflict = await client.post(f"/api/documents/{document_id}/reindex", headers=header)
    assert processing_conflict.status_code == 409

    deleted = await client.delete(f"/api/documents/{document_id}", headers=header)
    assert deleted.status_code == 200
    with SessionLocal() as db:
        assert db.get(Document, document_id) is None


async def test_failed_document_can_be_requeued(client: AsyncClient, monkeypatch) -> None:
    header = create_admin_header()
    monkeypatch.setattr(document_service, "submit", lambda _document_id, _job_id: None)
    uploaded = await client.post(
        "/api/documents",
        headers=header,
        files={"file": ("retry.md", b"# Retry\n\ncontent", "text/markdown")},
    )
    document_id = uploaded.json()["data"]["document"]["id"]
    with SessionLocal() as db:
        document = db.get(Document, document_id)
        assert document is not None
        document.status = DocumentStatus.FAILED
        document.error = "fake failure"
        db.commit()

    response = await client.post(f"/api/documents/{document_id}/reindex", headers=header)

    assert response.status_code == 202
    with SessionLocal() as db:
        document = db.get(Document, document_id)
        assert document is not None
        assert document.status == DocumentStatus.QUEUED
        assert document.error is None


async def test_preview_docx_pdf_cache_and_missing_download(
    client: AsyncClient, monkeypatch
) -> None:
    header = create_admin_header()
    monkeypatch.setattr(document_service, "submit", lambda _document_id, _job_id: None)

    docx = DocxDocument()
    docx.add_heading("河海大学预览", level=1)
    docx.add_paragraph("这是 DOCX 原文提取测试。")
    docx_buffer = BytesIO()
    docx.save(docx_buffer)
    uploaded_docx = await client.post(
        "/api/documents",
        headers=header,
        files={
            "file": (
                "preview.docx",
                docx_buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    docx_id = uploaded_docx.json()["data"]["document"]["id"]
    docx_preview = await client.get(f"/api/documents/{docx_id}/preview", headers=header)
    assert docx_preview.status_code == 200
    assert "DOCX 原文提取测试" in docx_preview.json()["data"]["content"]
    assert docx_preview.json()["data"]["format"] == "docx extracted-text"

    monkeypatch.setattr(
        "app.services.documents.parse_file",
        lambda _path, fallback_title: ParsedDocument(
            fallback_title,
            [ParsedSection("这是 PDF 第一页文本。", page_number=1)],
        ),
    )
    uploaded_pdf = await client.post(
        "/api/documents",
        headers=header,
        files={"file": ("preview.pdf", b"%PDF-test", "application/pdf")},
    )
    pdf_id = uploaded_pdf.json()["data"]["document"]["id"]
    pdf_preview = await client.get(f"/api/documents/{pdf_id}/preview", headers=header)
    assert pdf_preview.status_code == 200
    assert "[第 1 页]" in pdf_preview.json()["data"]["content"]
    assert "PDF 第一页文本" in pdf_preview.json()["data"]["content"]

    with SessionLocal() as db:
        pdf_document = db.get(Document, pdf_id)
        assert pdf_document is not None
        source = settings.upload_dir / pdf_document.stored_name
    source.unlink()
    missing_download = await client.get(
        f"/api/documents/{pdf_id}/download", headers=header
    )
    assert missing_download.status_code == 404
